from docx import Document
from docx.text.paragraph import Paragraph
from .models import *
from django.shortcuts import render, redirect
import xml.etree.ElementTree as ET
from .views import *
import os.path
from datetime import datetime
import re
from docs.working_directory import working_directory


def create_docx(request,project_id):
    indexhead = "Generate Project Document"
    hidesearch = "hide"
    member = Member.objects.get(user=request.user)
    project = Project.objects.get(id=project_id)
    
    if request.method == "POST":

        magic = '''<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Transitional//EN"
            "http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd" [
            <!ENTITY nbsp ' '>
            ]>''' 

        cleanr = re.compile('<.*?>|&([a-z0-9]+|#[0-9]{1,6}|#x[0-9a-f]{1,6});')
        document = Document()
        document.add_heading(project.project_title, 0)
        p = document.add_paragraph(re.sub(cleanr, '',  project.description))

        # Then Adding Viewpoints
        viewpoints = enumerate(Viewpoint.objects.filter(project=project, status="accepted"), start=1)
        if Viewpoint.objects.filter(project=project).exists():
             document.add_heading(project.project_title + '` Viewpoints', level=2)
        for num, viewpoint in viewpoints:
           document.add_heading(str(num)+': ' + viewpoint.viewpoint_name, level=3)
           document.add_paragraph(re.sub(cleanr, '', viewpoint.description))
        # getting project goals
        
        goals = enumerate(ViewpointGoal.objects.filter(goal__project=project, goal__status="accepted"), start=1)
        if Goal.objects.filter(project=project).exists():
            document.add_heading(project.project_title+'` Goals', level=2)
        for num, goal in goals:
            document.add_heading(str(num)+': ' + goal.goal.goal_name, level=3)
            document.add_paragraph('A Goal from ' + str(goal.viewpoint))
            document.add_paragraph(re.sub(cleanr, '',  goal.goal.description))


        # getting requirements
        requirements = enumerate(Requirement.objects.filter(project=project,status="accepted"), start=1)
        if Requirement.objects.filter(project=project).exists():
            document.add_heading(project.project_title+'` Requirements', level=2)
        for num, requirement in requirements:
            document.add_heading(str(num)+': ' + requirement.name, level=3)
            document.add_paragraph(re.sub(cleanr, '', requirement.description))
        
        # getting scenarios
            
            if RequirementScenario.objects.filter(requirement=requirement,scenario__status="accepted").exists():
                document.add_heading(str(requirement.name) +'` Scenarios', level=2)
                scenarios = enumerate(RequirementScenario.objects.filter(requirement=requirement, scenario__status="accepted"), start=1)
                for num, scenario in scenarios:
                    document.add_heading(str(num)+': ' + scenario.scenario.name, level=3)
                    document.add_paragraph(re.sub(cleanr, '', scenario.scenario.description))
            
            
            if RequirementProcess.objects.filter(requirement=requirement,process__status="accepted").exists():
                document.add_heading(str(requirement.name) + '` Processes', level=2)
                # getting processes
                processes = enumerate(RequirementProcess.objects.filter(requirement=requirement, process__status="accepted"), start=1)
                for num, process in processes:
                    document.add_heading(str(num)+': ' + process.process.process_name, level=3)
                    document.add_paragraph(re.sub(cleanr, '', process.process.description))
        now = datetime.now()
        number_of_generated_docs = GeneratedDocs.objects.filter(project=project).count()
        version = number_of_generated_docs + 1 
        
        filename = str(now)+project.project_title+'-Version-' + str(version) + '.docx'
        
        path = working_directory + '/docs/generated_docs/'+filename
        document.save(path)
        notification(request)
        create_document = GeneratedDocs.objects.create(
            project=project,docx=filename
        )
        create_document.save()
    docs = enumerate(GeneratedDocs.objects.filter(project=project).order_by('-id'), start=1)
    return render(request,'projects/my_projects/generate_docs.html',{
        'project':project,
        'project_id':project.id,
        'notification':notification,
        'indexhead':indexhead,
        'hidesearch':hidesearch,
        'docs':docs,
        'member':member
    })

def generated_docs(request,project_id):
    project = Project.objects.get(id=project_id)
    member = Member.objects.get(user=request.user)
    docs = enumerate(GeneratedDocs.objects.filter(project=project).order_by('-id'), start=1)
    return render(request,'projects/other_projects/project_generated_docs.html',{
        'project':project,
        'project_id':project.id,
        'notification':notification,
        'indexhead':"Project Generated Document(S)",
        'hidesearch':1,
        'docs':docs,
        'member':member
    })



def delete_file(request,file_id):
    file = GeneratedDocs.objects.get(id=file_id)
    # deleteing file from the drive
    os.remove(working_directory + '/docs/generated_docs/' + file.docx )
    delete = GeneratedDocs.objects.filter(id=file_id).delete()
    return redirect('projects:generatedocs',project_id=file.project.id)
